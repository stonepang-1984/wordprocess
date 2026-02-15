'''
The MIT License (MIT)

Copyright (c) 2026 Xiaofeng Pang woshi1984lian@163.com

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
'''

# coding: utf-8

import os
import re
from docx.oxml.table import *
from docx.oxml.text import paragraph
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.table import *
from docx.oxml.text import paragraph
from docx.shared import Inches, Pt, RGBColor
from io import BytesIO
f = BytesIO()


def Write_word_simple(word_name):
    '''解析docr，提取索引'''
    docr = Document(word_name)  # 读取的文件对象
    docw = Document()  # 需要写入的对象
    docw.styles['Normal'].font.name = u'宋体'
    docw.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    pindex = 0
    image_index = 0
    for i, p in enumerate(docr._body._element):
        if isinstance(p, CT_Tbl):
            docw._body._element._insert_tbl(p)
        if isinstance(p, paragraph.CT_P):
            stt = docr.paragraphs[pindex].text                              
            runs = docr.paragraphs[pindex].runs  # add_heading
            paragraph_ = docw.add_paragraph()  # 添加段落
            set_paragraph_style(docr, paragraph_, pindex)  # 设置段落样式
            for r in runs:
                style = get_run_style(r)
                rw = paragraph_.add_run(r.text)  # 添加run对象
                rw.font.name = r.font.name
                rw.font.size = Pt(12)
                if style.get('fontcolor'):
                    rw.font.color.rgb = style.get('fontcolor')
                rw.font.shadow = style.get('fontshadow')
                rw.font.strike = style.get('fontstrike')
                # print()
                rw.font.bold = style.get('fontbold')
                rw.font.italic = style.get('fontitalic')
                rw.font.underline = style.get('fontunderline')
                if r.style and r.style.name and r.style.name == "特别强调":
                    rw.font.bold = True
                    rw.bold = True
            if "rId" in p.xml:
                rId = re.findall('rId\d+', p.xml)[0]
                imgflag = save_image(docr,rId)
                if imgflag == True:
                    image_path = f"{image_index}.png"
                    docw.add_paragraph().add_run().add_picture(image_path,width=Inches(6))  # 添加图片
            pindex += 1
    ###表格样式调整
    for tb in docw.tables:
        tb.autofit=True
        tb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tb._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
        for row in tb.rows:
            for cell in row.cells:
                cell._tc.tcPr.tcW.type = 'auto'
                cell._tc.tcPr.tcW.w = 0
                for p in cell.paragraphs:
                    space_after = p.paragraph_format.space_after 
                    p.paragraph_format.space_before = space_after
    word_name_new = 'myhtml_new.docx'                
    docw.save(word_name_new)
    return word_name_new



def save_image(docr, rId):  # 保存外链图片到本地
    err_msg = 'If you want Embedding image function,please contact woshi1984lian@163.com'
    print(err_msg)


def get_run_style(run):
    '''获取run的样式'''
    item = {}
    item['fontsize'] = run.font.size
    item['fontcolor'] = run.font.color.rgb
    item['fontshadow'] = run.font.shadow
    item['fontstrike'] = run.font.strike  # 删除线
    item['fontitalic'] = run.font.italic  # 斜体
    item['fontunderline'] = run.font.underline  # 字体下划线
    item['fontbold'] = item['fontunderline']  #有下划线就加粗
    return item


def set_paragraph_style(docr, paragraph_, pindex):
    '''写入段落样式'''
    fl_indent = docr.paragraphs[pindex].paragraph_format.first_line_indent
    paragraph_.paragraph_format.first_line_indent = fl_indent  # 首行缩进
    line_spacing = docr.paragraphs[pindex].paragraph_format.line_spacing
    paragraph_.paragraph_format.line_spacing = line_spacing  # 行间距
    space_after = docr.paragraphs[pindex].paragraph_format.space_after
    paragraph_.paragraph_format.space_after = space_after  # 段落后间距
    space_before = docr.paragraphs[pindex].paragraph_format.space_before
    paragraph_.paragraph_format.space_before = space_before  # 段落后间距
    ###是否居中
    paragraph_.alignment = docr.paragraphs[pindex].paragraph_format.alignment


if __name__ == "__main__":
    htmlpath = 'myhtml.html'
    os.system(f'libreoffice --convert-to docx:"MS Word 2007 XML" {htmlpath}')
    word_name = "myhtml.docx"
    word_name_new = Write_word_simple(word_name)        

